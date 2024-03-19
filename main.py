from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.shared import Pt
from io import BytesIO
from PIL import Image
import base64
from xml.etree import ElementTree as ET
from docx.oxml.ns import qn

def add_img_to_cc(docx_content, img_dict):
 
 doc = Document(BytesIO(docx_content))
 pics_dic = {}  
 for cc_name, img_data in img_dict.items():
        # Decode base64 image data
    img_data_decoded = base64.b64decode(img_data.split(',')[1])
    
        # Add image to the document
    img_stream = BytesIO(img_data_decoded)
    # adding the image and removing it to create a rel between the document and the image
    doc.add_picture(img_stream)
    p = doc.paragraphs[-1]._element
    p.getparent().remove(p)
    
    # finding the id of the image
 
    keys = list(doc.part.rels.keys())
    
    # the rid of the picture we wanna add its also the last key in the list
   # search for the image content control
   #rid= keys[-1]
    pics_dic[cc_name]= keys[-1]   
 pic_ccs = doc._element.xpath("//a:blip")
     # iterating over the list of all the picture content controls
 i = 0 
 for pic in pic_ccs:
      real_name = pic_ccs[0].xpath("//pic:cNvPr")[i].name
      for name,rid in pics_dic.items():
        if(name == real_name ):  
          pic.set(qn("r:embed"), rid)
          try:
            temp = doc._element.xpath("//w:sdtPr")[0]
            temp.getparent().remove(temp)
          except:
            print("already gone")         
      i = i +1 
   # getting the content                 
 modified_docx_buffer = BytesIO()
 doc.save(modified_docx_buffer)
 modified_docx_content = modified_docx_buffer.getvalue()                 
 return modified_docx_content
def process_docx(fields:dict, docx_content:bytes) -> bytes:
  
    doc = Document(BytesIO(docx_content))

    # Iterate through each content control in the document
    for paragraph in doc.paragraphs:
        for cc in paragraph._element.xpath("w:sdt/w:sdtContent/w:r/w:t"):
            # Check each field for a match
            for key,value in fields.items():
                if cc.text == f"[{key}]":
                    cc.text = value

    modified_docx_buffer = BytesIO()
    doc.save(modified_docx_buffer)
    modified_docx_content = modified_docx_buffer.getvalue()                 
    return modified_docx_content

