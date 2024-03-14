from docx import Document
from docx.oxml.ns import qn

# function to replace the picture


def add_img_to_cc(file_name, img_path):

    doc = Document(file_name)

    # adding the image and removing it to create a rel between the document and the image
    doc.add_picture(img_path)
    p = doc.paragraphs[-1]._element
    p.getparent().remove(p)

    # finding the id of the image
    keys = list(doc.part.rels.keys())

    # the rid of the picture we wanna add its also the last key in the list
    rid = keys[-1]

    # search for the image content control
    pic_ccs = doc._element.xpath("//a:blip")

    # iterating over the list of all the picture content controls
    for pic in pic_ccs:
        pic.set(qn("r:embed"), rid)
        try:
            temp = doc._element.xpath("//w:sdtPr")[0]
            temp.getparent().remove(temp)
        except:
            print("already gone")

    return doc
